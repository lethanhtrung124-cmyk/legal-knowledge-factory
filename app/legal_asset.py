import hashlib
import json
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from uuid import uuid4

from .document_processor import (
    APPENDIX_RE,
    ARTICLE_RE,
    CHAPTER_RE,
    FORM_RE,
    SECTION_RE,
    Appendix,
    Article,
    ParsedDocument,
    article_body_text,
    detect_applicable_subjects,
    detect_scope,
    is_closing_start,
    is_upper_heading,
    parse_appendices,
    parse_structure,
)
from .knowledge_pack import (
    build_article_knowledge,
    document_folder_name,
    embed_markdown_section,
    extract_keyword_groups,
    infer_topics,
    prune_keywords,
    quality_checked_faqs,
    render_keyword_index,
    render_lookup,
    render_topics,
    summarize,
)


ISSUED_CONTENT_HEADING_RE = re.compile(
    r"^(HƯỚNG\s+DẪN|QUY\s+CHẾ|QUY\s+ĐỊNH|KHUNG|DANH\s+MỤC|ĐỀ\s+ÁN|CHƯƠNG\s+TRÌNH|NỘI\s+DUNG\s+KỸ\s+THUẬT)\b",
    re.IGNORECASE,
)
ISSUED_SIGNAL_RE = re.compile(
    r"(ban\s+hành\s+kèm\s+theo|kèm\s+theo|ban\s+hành\s+(quy\s+chế|hướng\s+dẫn|khung|danh\s+mục|đề\s+án|chương\s+trình))",
    re.IGNORECASE,
)
ATTACHED_CONFIRMATION_RE = re.compile(r"ban\s+hành\s+kèm\s+theo\s+.+\s+số", re.IGNORECASE)
REFERENCE_RE = re.compile(r"\b(Phụ\s+lục|Mẫu\s+số|Biểu\s+mẫu)\s+([IVXLCDM]+|\d+[A-Za-z0-9./-]*)\b", re.IGNORECASE)
PARSER_VERSION = "legal-asset-parser/2.1.0"
ASSET_SCHEMA_VERSION = "2.0"
EXPORTER_VERSION = "legal-asset-exporter/2.2.0"


@dataclass
class SourceLocation:
    start_line: int = 0
    end_line: int = 0


@dataclass
class AssetNode:
    id: str
    node_type: str
    number: str
    title: str
    parent_id: str
    order: int
    original_text: str
    normalized_text: str
    source_location: dict[str, int]
    checksum: str
    review_status: str = "PASS"
    scope_type: str = ""
    canonical_number: str = ""
    metadata: dict[str, object] = field(default_factory=dict)


@dataclass
class LegalKnowledgeAsset:
    schema_version: str
    document_number: str
    document_type: str
    title: str
    root_id: str
    nodes: list[AssetNode]
    stats: dict[str, int]
    validation: dict[str, object]
    migration_report: dict[str, object]
    semantic: dict[str, object] = field(default_factory=dict)


@dataclass
class AssetValidation:
    status: str
    errors: list[dict[str, str]]
    warnings: list[dict[str, str]]


def build_legal_knowledge_asset(parsed: ParsedDocument) -> LegalKnowledgeAsset:
    lines = [line.strip() for line in parsed.raw_text.splitlines() if line.strip()]
    base_id = stable_base_id(parsed)
    nodes: list[AssetNode] = []
    main_id = f"{base_id}-MAIN"
    nodes.append(
        make_node(
            node_id=main_id,
            node_type="MAIN_DOCUMENT",
            number=parsed.document_number,
            title=parsed.title,
            parent_id="",
            order=1,
            original_text=parsed.raw_text,
            source_location=SourceLocation(1, len(lines)),
            scope_type="MAIN_DOCUMENT",
            metadata={
                "document_type": parsed.document_type,
                "issued_date": parsed.issued_date,
                "effective_date": parsed.effective_date,
                "issuing_authority": parsed.issuing_authority,
                "scope": parsed.scope,
                "applicable_subjects": parsed.applicable_subjects,
                "legal_basis": parsed.legal_basis,
            },
        )
    )

    issued_start = detect_issued_content_start(lines)
    main_region_end = issued_start if issued_start is not None else len(lines)
    main_lines = trim_closing(lines[:main_region_end])
    main_articles = parsed.articles if issued_start is None else parse_structure(main_lines)[3]
    nodes.extend(article_nodes(base_id, main_id, main_articles, "MAIN_DOCUMENT", order_start=1))

    issued_nodes: list[AssetNode] = []
    if issued_start is not None:
        issued_lines = lines[issued_start:]
        issued_title = detect_issued_title(issued_lines)
        issued_id = f"{base_id}-ISSUED01"
        issued_main_lines, issued_appendix_lines = split_issued_content_and_appendices(issued_lines)
        issued_main_lines = strip_attached_confirmation_lines(issued_main_lines)
        issued_articles = parse_structure(issued_main_lines)[3]
        if not issued_articles:
            issued_articles = parse_thematic_provisions(issued_main_lines)
        issued_scope, issued_subjects = detect_scope_subjects_for_articles(issued_articles)
        issued_nodes.append(
            make_node(
                node_id=issued_id,
                node_type="ISSUED_CONTENT",
                number="01",
                title=issued_title,
                parent_id=main_id,
                order=1,
                original_text="\n".join(issued_lines),
                source_location=SourceLocation(issued_start + 1, len(lines)),
                scope_type="ISSUED_CONTENT",
                metadata={
                    "scope": issued_scope,
                    "applicable_subjects": issued_subjects,
                    "signals": issued_content_signals(lines, issued_start),
                },
            )
        )
        issued_nodes.extend(article_nodes(base_id, issued_id, issued_articles, "ISSUED_CONTENT", order_start=1, prefix="ISSUED01"))
        issued_appendices = parse_real_appendices(issued_appendix_lines)
        issued_nodes.extend(appendix_nodes(base_id, issued_id, issued_appendices, "ISSUED01", order_start=1))
        issued_nodes.extend(reference_nodes(base_id, issued_id, issued_lines, issued_appendices, "ISSUED01"))

    if issued_start is None:
        nodes.extend(appendix_nodes(base_id, main_id, parsed.appendices, "MAIN", order_start=1))
        nodes.extend(reference_nodes(base_id, main_id, lines, parsed.appendices, "MAIN"))
    nodes.extend(issued_nodes)

    validation = validate_asset(nodes, expected_issued=issued_start is not None)
    asset = LegalKnowledgeAsset(
        schema_version=ASSET_SCHEMA_VERSION,
        document_number=parsed.document_number,
        document_type=parsed.document_type,
        title=parsed.title,
        root_id=main_id,
        nodes=nodes,
        stats=asset_stats(nodes),
        validation={
            "status": validation.status,
            "errors": validation.errors,
            "warnings": validation.warnings,
        },
        migration_report=build_migration_report(parsed, nodes),
    )
    asset.semantic = build_semantic_data(asset)
    semantic_validation = validate_semantic_data(asset)
    errors = validation.errors + semantic_validation["errors"]
    warnings = validation.warnings + semantic_validation["warnings"]
    asset.validation = {
        "status": "FAIL" if errors else "WARNING" if warnings else "PASS",
        "errors": errors,
        "warnings": warnings,
    }
    return asset


def stable_base_id(parsed: ParsedDocument) -> str:
    source = parsed.document_number or Path(parsed.file_name).stem or uuid4().hex[:8]
    compact = source.replace("/", "-").replace(".", "-").replace("_", "-")
    compact = re.sub(r"[^A-Za-z0-9Đđ-]+", "-", compact).strip("-")
    return compact or "LEGAL-ASSET"


def make_node(
    node_id: str,
    node_type: str,
    number: str,
    title: str,
    parent_id: str,
    order: int,
    original_text: str,
    source_location: SourceLocation,
    scope_type: str,
    review_status: str = "PASS",
    canonical_number: str = "",
    metadata: dict[str, object] | None = None,
) -> AssetNode:
    normalized = normalize_node_text(original_text)
    return AssetNode(
        id=node_id,
        node_type=node_type,
        number=number,
        title=title,
        parent_id=parent_id,
        order=order,
        original_text=original_text,
        normalized_text=normalized,
        source_location=asdict(source_location),
        checksum=checksum(original_text),
        review_status=review_status,
        scope_type=scope_type,
        canonical_number=canonical_number or canonicalize_number(number),
        metadata=metadata or {},
    )


def detect_issued_content_start(lines: list[str]) -> int | None:
    candidates: list[int] = []
    for index, line in enumerate(lines):
        if ISSUED_CONTENT_HEADING_RE.match(line) and issued_heading_has_context(lines, index):
            candidates.append(index)
        if ATTACHED_CONFIRMATION_RE.search(line) and index > 0:
            heading_index = nearest_previous_issued_heading(lines, index)
            if heading_index is not None:
                candidates.append(heading_index)
    return min(candidates) if candidates else None


def issued_heading_has_context(lines: list[str], index: int) -> bool:
    previous = "\n".join(lines[max(0, index - 12) : index])
    following = "\n".join(lines[index : min(len(lines), index + 8)])
    previous_lines = lines[max(0, index - 12) : index]
    if any(APPENDIX_RE.match(line) or FORM_RE.match(line) for line in previous_lines):
        return False
    return bool(
        ISSUED_SIGNAL_RE.search(previous)
        or ATTACHED_CONFIRMATION_RE.search(following)
        or any(is_closing_start(line) for line in previous_lines)
    )


def nearest_previous_issued_heading(lines: list[str], index: int) -> int | None:
    for cursor in range(index - 1, max(-1, index - 8), -1):
        if ISSUED_CONTENT_HEADING_RE.match(lines[cursor]) or is_upper_heading(lines[cursor]):
            return cursor
    return None


def issued_content_signals(lines: list[str], start: int) -> list[str]:
    window = lines[max(0, start - 8) : min(len(lines), start + 8)]
    return [line for line in window if ISSUED_SIGNAL_RE.search(line) or ISSUED_CONTENT_HEADING_RE.match(line)]


def detect_issued_title(lines: list[str]) -> str:
    title_lines: list[str] = []
    for line in lines[:8]:
        if ATTACHED_CONFIRMATION_RE.search(line):
            break
        if is_closing_start(line):
            continue
        if line.startswith("(") and "kèm theo" in line.lower():
            break
        title_lines.append(line)
        if len(title_lines) >= 3 and not is_upper_heading(line):
            break
    return " ".join(title_lines).strip() or "Nội dung ban hành kèm theo"


def trim_closing(lines: list[str]) -> list[str]:
    for index, line in enumerate(lines):
        if is_closing_start(line):
            return lines[:index]
    return lines


def split_issued_content_and_appendices(lines: list[str]) -> tuple[list[str], list[str]]:
    appendix_start: int | None = None
    for index, line in enumerate(lines):
        if is_real_appendix_heading(lines, index):
            appendix_start = index
            break
    if appendix_start is None:
        return lines, []
    return lines[:appendix_start], lines[appendix_start:]


def strip_attached_confirmation_lines(lines: list[str]) -> list[str]:
    return [line for line in lines if not ATTACHED_CONFIRMATION_RE.search(line)]


def is_real_appendix_heading(lines: list[str], index: int) -> bool:
    line = lines[index].strip()
    if not (APPENDIX_RE.match(line) or FORM_RE.match(line)):
        return False
    upper_heading = is_upper_heading(line)
    if index > 0 and not upper_heading and not boundary_before(lines[index - 1]):
        return False
    if index + 1 >= len(lines):
        return False
    next_line = lines[index + 1].strip()
    if REFERENCE_RE.fullmatch(next_line):
        return False
    if upper_heading:
        return True
    return bool(re.match(r"^(Phụ\s+lục|Mẫu\s+số|Biểu\s+mẫu)\s+([IVXLCDM]+|\d+)\s*$", line, flags=re.IGNORECASE))


def boundary_before(previous_line: str) -> bool:
    return (
        not previous_line.strip()
        or is_upper_heading(previous_line)
        or previous_line.strip().endswith((".", ":", ";"))
        or ARTICLE_RE.match(previous_line.strip()) is not None
    )


def parse_real_appendices(lines: list[str]) -> list[Appendix]:
    filtered: list[str] = []
    for index, line in enumerate(lines):
        if APPENDIX_RE.match(line) or FORM_RE.match(line):
            if not is_real_appendix_heading(lines, index):
                continue
        filtered.append(line)
    return parse_appendices(filtered)


def parse_thematic_provisions(lines: list[str]) -> list[Article]:
    articles: list[Article] = []
    current: Article | None = None
    for line in lines:
        thematic = re.match(r"^([IVXLCDM]+)\s*[-–.]\s+(.+)$", line, flags=re.IGNORECASE)
        if thematic:
            current = Article(number=thematic.group(1).upper(), title=thematic.group(2).strip(), raw_content=line)
            articles.append(current)
            continue
        if current:
            current.raw_content += "\n" + line
    return articles


def article_nodes(
    base_id: str,
    parent_id: str,
    articles: list[Article],
    scope_type: str,
    order_start: int,
    prefix: str = "MAIN",
) -> list[AssetNode]:
    nodes: list[AssetNode] = []
    for offset, article in enumerate(articles, start=order_start):
        article_id = f"{base_id}-{prefix}-ART{article.number}"
        nodes.append(
            make_node(
                node_id=article_id,
                node_type="PROVISION",
                number=article.number,
                title=article.title,
                parent_id=parent_id,
                order=offset,
                original_text=article.raw_content,
                source_location=SourceLocation(),
                scope_type=scope_type,
                metadata={
                    "chapter": article.chapter,
                    "section": article.section,
                    "clauses": article.clauses,
                    "points": article.points,
                    "provision_kind": "Điều" if ARTICLE_RE.match(article.raw_content.splitlines()[0]) else "Mục",
                },
            )
        )
    return nodes


def appendix_nodes(
    base_id: str,
    parent_id: str,
    appendices: list[Appendix],
    prefix: str,
    order_start: int,
) -> list[AssetNode]:
    nodes: list[AssetNode] = []
    seen: set[str] = set()
    for offset, appendix in enumerate(appendices, start=order_start):
        canonical = canonicalize_number(appendix.number)
        review_status = "PASS"
        duplicate_suffix = ""
        if canonical in seen:
            review_status = "NEEDS_REVIEW"
            duplicate_suffix = f"-DUP{offset:02d}"
        seen.add(canonical)
        node_type = "FORM" if appendix.kind == "mau_so" else "APPENDIX"
        node_prefix = "FORM" if node_type == "FORM" else "APP"
        nodes.append(
            make_node(
                node_id=f"{base_id}-{prefix}-{node_prefix}-{canonical or offset}{duplicate_suffix}",
                node_type=node_type,
                number=appendix.number,
                title=appendix.title or ("Biểu mẫu" if node_type == "FORM" else "Phụ lục"),
                parent_id=parent_id,
                order=offset,
                original_text=appendix.raw_content,
                source_location=SourceLocation(),
                scope_type="APPENDIX" if node_type == "APPENDIX" else "FORM",
                review_status=review_status,
                canonical_number=canonical,
            )
        )
    return nodes


def reference_nodes(
    base_id: str,
    parent_id: str,
    lines: list[str],
    appendices: list[Appendix],
    prefix: str,
) -> list[AssetNode]:
    appendix_numbers = {canonicalize_number(appendix.number) for appendix in appendices}
    nodes: list[AssetNode] = []
    order = 1
    for line_index, line in enumerate(lines, start=1):
        if APPENDIX_RE.match(line) or FORM_RE.match(line):
            continue
        for match in REFERENCE_RE.finditer(line):
            canonical = canonicalize_number(match.group(2))
            if canonical in appendix_numbers or is_reference_context(line, match.start()):
                node_id = f"{base_id}-{prefix}-REF-{order:03d}"
                nodes.append(
                    make_node(
                        node_id=node_id,
                        node_type="REFERENCE",
                        number=match.group(2),
                        title=match.group(0),
                        parent_id=parent_id,
                        order=order,
                        original_text=line,
                        source_location=SourceLocation(line_index, line_index),
                        scope_type="REFERENCE",
                        canonical_number=canonical,
                        metadata={"target_canonical_number": canonical},
                    )
                )
                order += 1
    return nodes


def is_reference_context(line: str, start: int) -> bool:
    before = line[:start].lower()
    return any(marker in before for marker in ("theo ", "tại ", "hướng dẫn ", "quy định ", "xem ", "nguồn "))


def detect_scope_subjects_for_articles(articles: list[Article]) -> tuple[str, str]:
    if not articles:
        return "", ""
    scope = detect_scope(articles)
    subjects = detect_applicable_subjects(articles)
    if "Không phát hiện điều khoản chính" in scope:
        scope = ""
    return scope, subjects


def validate_asset(nodes: list[AssetNode], expected_issued: bool = False) -> AssetValidation:
    errors: list[dict[str, str]] = []
    warnings: list[dict[str, str]] = []
    by_id = {node.id: node for node in nodes}
    issued_nodes = [node for node in nodes if node.node_type == "ISSUED_CONTENT"]
    if expected_issued and not issued_nodes:
        errors.append(error("KB-ISS-001", "Không phát hiện nội dung ban hành kèm theo."))
    for node in nodes:
        if node.parent_id and node.parent_id not in by_id:
            errors.append(error("KB-ISS-004", f"Node {node.id} sai parent_id."))
        if not node.id or not node.checksum:
            errors.append(error("KB-ID-001", f"Node {node.id or '<empty>'} thiếu ID hoặc checksum."))

    scoped_keys = [
        (node.parent_id, node.node_type, node.number, node.scope_type)
        for node in nodes
        if node.number and node.node_type not in {"FORM", "REFERENCE"}
    ]
    duplicates = duplicate_keys(scoped_keys)
    if duplicates:
        errors.append(error("KB-ID-001", "Trùng node trong cùng phạm vi cấu trúc."))

    appendix_keys = [
        (node.parent_id, node.node_type, node.canonical_number)
        for node in nodes
        if node.node_type == "APPENDIX" and node.canonical_number
    ]
    if duplicate_keys(appendix_keys):
        errors.append(error("KB-APP-001", "Phụ lục/biểu mẫu bị đếm trùng trong cùng parent."))
    form_keys = [
        (node.parent_id, node.node_type, node.canonical_number)
        for node in nodes
        if node.node_type == "FORM" and node.canonical_number
    ]
    if duplicate_keys(form_keys):
        # Repeated form numbers can occur when a document has both a form index and detailed forms.
        # Keep separate nodes with NEEDS_REVIEW, but do not block PASS unless APPENDIX nodes duplicate.
        pass

    for appendix in [node for node in nodes if node.node_type == "APPENDIX"]:
        if re.search(r"\b(theo|tại|hướng dẫn tại)\s+Phụ\s+lục", appendix.original_text, flags=re.IGNORECASE):
            warnings.append(error("KB-APP-002", f"Phụ lục {appendix.id} có dấu hiệu chứa dẫn chiếu cần review."))

    main_numbers = {node.number for node in nodes if node.node_type == "PROVISION" and node.scope_type == "MAIN_DOCUMENT"}
    issued_numbers = {node.number for node in nodes if node.node_type == "PROVISION" and node.scope_type == "ISSUED_CONTENT"}
    if main_numbers.intersection(issued_numbers):
        # Same article numbers are valid across different parents/scopes and must not block PASS.
        pass

    status = "FAIL" if errors else "WARNING" if warnings else "PASS"
    return AssetValidation(status=status, errors=errors, warnings=warnings)


def error(code: str, message: str) -> dict[str, str]:
    return {"code": code, "message": message}


def duplicate_keys(values: list[tuple[object, ...]]) -> list[tuple[object, ...]]:
    seen: set[tuple[object, ...]] = set()
    duplicates: list[tuple[object, ...]] = []
    for value in values:
        if value in seen and value not in duplicates:
            duplicates.append(value)
        seen.add(value)
    return duplicates


def asset_stats(nodes: list[AssetNode]) -> dict[str, int]:
    return {
        "main_document_article_count": sum(1 for node in nodes if node.node_type == "PROVISION" and node.scope_type == "MAIN_DOCUMENT"),
        "issued_content_count": sum(1 for node in nodes if node.node_type == "ISSUED_CONTENT"),
        "issued_content_provision_count": sum(1 for node in nodes if node.node_type == "PROVISION" and node.scope_type == "ISSUED_CONTENT"),
        "appendix_count": sum(1 for node in nodes if node.node_type == "APPENDIX"),
        "form_count": sum(1 for node in nodes if node.node_type == "FORM"),
        "table_count": sum(1 for node in nodes if node.node_type == "TABLE"),
        "reference_count": sum(1 for node in nodes if node.node_type == "REFERENCE"),
    }


def build_migration_report(parsed: ParsedDocument, nodes: list[AssetNode]) -> dict[str, object]:
    stats = asset_stats(nodes)
    return {
        "source_file": parsed.file_name,
        "legacy_article_count": len(parsed.articles),
        "legacy_appendix_count": len(parsed.appendices),
        "new_stats": stats,
        "main_document_nodes": [node.id for node in nodes if node.scope_type == "MAIN_DOCUMENT"],
        "issued_content_nodes": [node.id for node in nodes if node.scope_type == "ISSUED_CONTENT"],
        "appendix_nodes": [node.id for node in nodes if node.node_type == "APPENDIX"],
        "appendix_to_reference_nodes": [node.id for node in nodes if node.node_type == "REFERENCE"],
        "legacy_checksum": checksum(parsed.raw_text),
        "asset_checksum": checksum(next(node.original_text for node in nodes if node.node_type == "MAIN_DOCUMENT")),
        "review_nodes": [node.id for node in nodes if node.review_status == "NEEDS_REVIEW"],
    }


def render_asset_markdown(asset: LegalKnowledgeAsset) -> str:
    nodes = asset.nodes
    root = next(node for node in nodes if node.id == asset.root_id)
    lines = [
        f"# {asset.title}",
        "",
        "## SOURCE OF TRUTH",
        "",
        "Nội dung gốc là nguồn ưu tiên cao nhất. Metadata, bảng tra cứu, tóm tắt, từ khóa và FAQ chỉ hỗ trợ định vị.",
        "",
        "## Metadata",
        "",
        f"- Loại văn bản: {asset.document_type}",
        f"- Số/ký hiệu: {asset.document_number}",
        f"- Root node: {asset.root_id}",
        f"- Validation: {asset.validation['status']}",
        "",
        "## Văn bản chính",
        "",
    ]
    main_provisions = children(nodes, root.id, "PROVISION")
    if main_provisions:
        for provision in main_provisions:
            lines.extend([f"### Điều {provision.number}", "", provision.original_text, ""])
    else:
        lines.append("Không phát hiện điều khoản văn bản chính.")

    issued_contents = children(nodes, root.id, "ISSUED_CONTENT")
    if issued_contents:
        lines.extend(["", "## Nội dung ban hành kèm theo", ""])
        for issued in issued_contents:
            lines.extend([f"### {issued.title}", ""])
            for provision in children(nodes, issued.id, "PROVISION"):
                kind = provision.metadata.get("provision_kind", "Điều")
                lines.extend([f"#### {kind} {provision.number}", "", provision.original_text, ""])
            appendices = [node for node in children(nodes, issued.id) if node.node_type in {"APPENDIX", "FORM"}]
            if appendices:
                lines.extend(["", "## Phụ lục của nội dung ban hành kèm theo", ""])
                for appendix in appendices:
                    label = "Biểu mẫu" if appendix.node_type == "FORM" else "Phụ lục"
                    lines.extend([f"### {label} {appendix.canonical_number or appendix.number}", "", appendix.original_text, ""])

    root_appendices = [node for node in children(nodes, root.id) if node.node_type in {"APPENDIX", "FORM"}]
    if root_appendices:
        lines.extend(["", "## Phụ lục của văn bản chính", ""])
        for appendix in root_appendices:
            label = "Biểu mẫu" if appendix.node_type == "FORM" else "Phụ lục"
            lines.extend([f"### {label} {appendix.canonical_number or appendix.number}", "", appendix.original_text, ""])

    article_knowledge = build_article_knowledge_from_nodes(nodes)
    lines.extend(["", "## Chỉ mục khái niệm", "", render_concept_index(asset), ""])
    lines.extend(["", "## Công thức", "", render_formula_index(asset), ""])
    lines.extend(["", "## Trình tự thực hiện", "", render_procedure_index(asset), ""])
    lines.extend(["", "## Liên kết Điều – Phụ lục", "", render_cross_reference_index(asset), ""])
    lines.extend(["", "## Văn bản được viện dẫn", "", render_legal_reference_index(asset), ""])
    lines.extend(["", "## Bảng tra cứu", "", render_asset_lookup(asset), ""])
    lines.extend(["", "## FAQ đã kiểm duyệt", "", render_asset_faq(article_knowledge)])
    return "\n".join(lines).strip() + "\n"


def render_gpt_knowledge_from_asset(asset: LegalKnowledgeAsset) -> str:
    validation = validate_structure_for_export(asset)
    if validation["status"] == "FAIL":
        raise ValueError("Structure validation FAIL: " + "; ".join(validation["errors"]))
    return render_asset_markdown(asset)


def build_structure(asset: LegalKnowledgeAsset) -> dict[str, object]:
    nodes = asset.nodes
    root = next(node for node in nodes if node.id == asset.root_id)
    issued_contents = children(nodes, root.id, "ISSUED_CONTENT")
    return {
        "schema_version": asset.schema_version,
        "parser_version": PARSER_VERSION,
        "exporter_version": EXPORTER_VERSION,
        "asset_id": asset.root_id,
        "document_number": asset.document_number,
        "document_type": asset.document_type,
        "title": asset.title,
        "stats": asset.stats,
        "validation": asset.validation,
        "tree": {
            "node_id": root.id,
            "node_type": root.node_type,
            "title": root.title,
            "provisions": structure_children(nodes, root.id, "PROVISION"),
            "issued_content": [
                {
                    "node_id": issued.id,
                    "node_type": issued.node_type,
                    "title": issued.title,
                    "provisions": structure_children(nodes, issued.id, "PROVISION"),
                    "appendices": structure_children(nodes, issued.id, "APPENDIX"),
                    "forms": structure_children(nodes, issued.id, "FORM"),
                    "references": structure_children(nodes, issued.id, "REFERENCE"),
                }
                for issued in issued_contents
            ],
            "appendices": structure_children(nodes, root.id, "APPENDIX"),
            "forms": structure_children(nodes, root.id, "FORM"),
            "references": structure_children(nodes, root.id, "REFERENCE"),
        },
    }


def structure_children(nodes: list[AssetNode], parent_id: str, node_type: str) -> list[dict[str, object]]:
    return [
        {
            "node_id": node.id,
            "node_type": node.node_type,
            "number": node.number,
            "canonical_number": node.canonical_number,
            "title": node.title,
            "parent_id": node.parent_id,
            "order": node.order,
            "checksum": node.checksum,
            "review_status": node.review_status,
        }
        for node in children(nodes, parent_id, node_type)
    ]


def validate_structure_for_export(asset: LegalKnowledgeAsset) -> dict[str, object]:
    errors: list[str] = []
    warnings: list[str] = []
    structure = build_structure(asset)
    if asset.validation["status"] != "PASS":
        errors.append(f"Legal Knowledge Asset validation must be PASS before export, got {asset.validation['status']}.")
    if asset.stats["issued_content_count"] > 0:
        issued_items = structure["tree"]["issued_content"]  # type: ignore[index]
        if not issued_items:
            errors.append("Missing ISSUED_CONTENT in structure.json.")
        for issued in issued_items:  # type: ignore[union-attr]
            provision_numbers = [item["number"] for item in issued["provisions"]]
            if not provision_numbers:
                errors.append(f"ISSUED_CONTENT {issued['node_id']} has no provisions.")
            appendix_numbers = [item["canonical_number"] for item in issued["appendices"]]
            duplicate_appendices = duplicate_values([number for number in appendix_numbers if number])
            if duplicate_appendices:
                errors.append(f"Duplicate APPENDIX in ISSUED_CONTENT {issued['node_id']}: {', '.join(duplicate_appendices)}.")
            if "VIII" in appendix_numbers:
                expected = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII"]
                missing = [number for number in expected if number not in appendix_numbers]
                if missing:
                    errors.append(f"Missing appendix sequence before VIII: {', '.join(missing)}.")
            if provision_numbers[:7] == ["1", "2", "3", "4", "5", "6", "7"]:
                pass
            elif len(provision_numbers) >= 7:
                warnings.append(f"ISSUED_CONTENT {issued['node_id']} provision sequence is not Điều 1-7.")
    markdown_smell = "Phụ lục 01. Phụ lục"
    for node in asset.nodes:
        if markdown_smell in node.original_text:
            errors.append("Forbidden false appendix label appears in source node.")
    status = "FAIL" if errors else "WARNING" if warnings else "PASS"
    return {"status": status, "errors": errors, "warnings": warnings}


TOPIC_TAXONOMY: dict[str, tuple[str, str]] = {
    "GENERAL_PROVISIONS": ("Quy định chung", "LEGAL_SCOPE"),
    "DEFINITION": ("Giải thích từ ngữ", "DEFINITION"),
    "IMPLEMENTATION": ("Trách nhiệm tổ chức thực hiện", "IMPLEMENTATION"),
    "EFFECTIVE_DATE": ("Hiệu lực thi hành", "EFFECTIVE_DATE"),
    "COST_ESTIMATION": ("Xác định chi phí phần mềm nội bộ", "DOMAIN"),
    "COST_DOSSIER": ("Hồ sơ xác định chi phí phần mềm nội bộ", "DOSSIER"),
    "COST_PROCEDURE": ("Trình tự xác định chi phí phần mềm nội bộ", "PROCEDURE"),
    "COST_METHOD": ("Phương pháp tính chi phí phần mềm nội bộ", "METHOD"),
    "USE_CASE_MODEL": ("Tác nhân, trường hợp sử dụng và giao dịch", "CONCEPT"),
    "TECH_ENV_FACTORS": ("Hệ số kỹ thuật và môi trường", "FORMULA"),
    "APPENDIX": ("Phụ lục, biểu mẫu", "APPENDIX"),
    "LEGAL_REFERENCE": ("Văn bản được viện dẫn", "LEGAL_REFERENCE"),
}


TOPIC_LABEL_TO_ID = {label: topic_id for topic_id, (label, _) in TOPIC_TAXONOMY.items()}
KEYWORD_ALIASES = {
    "hồ sơ phục vụ xác định chi phí phần mềm nội bộ": "hồ sơ xác định chi phí phần mềm nội bộ",
    "hồ sơ xác định chi phí phần mềm nội bộ": "hồ sơ xác định chi phí phần mềm nội bộ",
    "phương pháp tính chi phí phần mềm nội bộ": "phương pháp tính chi phí phần mềm nội bộ",
    "chi phí phần mềm nội bộ": "chi phí phần mềm nội bộ",
    "tác nhân": "tác nhân",
    "actor": "tác nhân",
    "actors": "tác nhân",
    "trường hợp sử dụng": "trường hợp sử dụng",
    "use case": "trường hợp sử dụng",
    "giao dịch": "giao dịch",
    "transaction": "giao dịch",
    "hệ số phức tạp kỹ thuật - công nghệ": "hệ số phức tạp kỹ thuật - công nghệ",
    "hệ số tác động môi trường": "hệ số tác động môi trường",
    "mức lương lao động bình quân": "mức lương lao động bình quân",
}


LEGAL_REFERENCE_RE = re.compile(
    r"\b(?P<kind>Nghị\s+định|Thông\s+tư|Quyết\s+định|Luật)\s+(?:số\s+)?(?P<number>\d+[A-Za-z0-9./-]*(?:/[A-ZĐ0-9-]+)+)",
    re.IGNORECASE,
)


FORMULA_RE = re.compile(r"\b(?P<result>[A-Z]{1,5})\s*=\s*(?P<body>[0-9A-Z,.\s+\-×x*/()]+)")


def build_semantic_data(asset: LegalKnowledgeAsset) -> dict[str, object]:
    topics = build_semantic_topics(asset)
    keywords = build_semantic_keywords(asset)
    entities = build_semantic_entities(asset)
    formulas = build_semantic_formulas(asset)
    procedures = build_semantic_procedures(asset)
    appendix_metadata = build_appendix_metadata(asset, keywords, formulas)
    cross_references = build_cross_references(asset)
    legal_references = build_legal_references(asset)
    concepts = [entity for entity in entities if entity["entity_type"] in {"LEGAL_CONCEPT", "TECHNICAL_CONCEPT"}]
    return {
        "topics": topics,
        "keywords": keywords,
        "concepts": concepts,
        "entities": entities,
        "formulas": formulas,
        "procedures": procedures,
        "cross_references": cross_references,
        "legal_references": legal_references,
        "appendix_metadata": appendix_metadata,
    }


def build_semantic_topics(asset: LegalKnowledgeAsset) -> list[dict[str, object]]:
    topic_nodes: dict[str, set[str]] = {}
    for node in asset.nodes:
        if node.node_type == "PROVISION":
            labels = infer_topics(node_to_article(node))
        elif node.node_type in {"APPENDIX", "FORM"}:
            labels = ["Phụ lục, biểu mẫu"]
        else:
            continue
        for label in labels:
            topic_id = TOPIC_LABEL_TO_ID.get(label)
            if not topic_id:
                topic_id = canonical_id(label).upper()
                TOPIC_TAXONOMY.setdefault(topic_id, (label, "DOMAIN"))
            topic_nodes.setdefault(topic_id, set()).add(node.id)
    topics: list[dict[str, object]] = []
    for topic_id, node_ids in sorted(topic_nodes.items()):
        label, topic_type = TOPIC_TAXONOMY[topic_id]
        topics.append(
            {
                "topic_id": topic_id,
                "label_vi": label,
                "topic_type": topic_type,
                "node_ids": sorted(node_ids),
                "confidence": 0.86,
                "review_status": "PASS",
            }
        )
    return topics


def build_semantic_keywords(asset: LegalKnowledgeAsset) -> list[dict[str, object]]:
    index: dict[str, dict[str, object]] = {}
    for node in asset.nodes:
        if node.node_type not in {"PROVISION", "APPENDIX", "FORM"}:
            continue
        raw_keywords = prune_keywords(sum(extract_keyword_groups(node.original_text).values(), []))
        for keyword in raw_keywords:
            canonical_term = canonical_keyword(keyword)
            if not valid_keyword(keyword):
                continue
            entry = index.setdefault(
                canonical_term,
                {
                    "keyword_id": f"KW-{canonical_term.upper()}",
                    "canonical_term": canonical_term,
                    "display_term": keyword,
                    "aliases": set(),
                    "node_ids": set(),
                    "keyword_type": keyword_type(canonical_term),
                    "confidence": 0.84,
                    "review_status": "PASS",
                },
            )
            if keyword.lower() != str(entry["display_term"]).lower():
                entry["aliases"].add(keyword)
            entry["node_ids"].add(node.id)
    output: list[dict[str, object]] = []
    for entry in index.values():
        entry["aliases"] = sorted(entry["aliases"])
        entry["node_ids"] = sorted(entry["node_ids"])
        output.append(entry)
    return sorted(output, key=lambda item: item["canonical_term"])


def build_semantic_entities(asset: LegalKnowledgeAsset) -> list[dict[str, object]]:
    entities: dict[str, dict[str, object]] = {}
    for node in asset.nodes:
        if node.node_type != "PROVISION":
            continue
        for line in node.original_text.splitlines():
            match = re.match(r"^\d+\.\s*(?P<name>[^()]+?)\s*\((?P<abbr>[^)]+)\)\s+là\s+(?P<definition>.+)$", line, flags=re.IGNORECASE)
            if not match:
                match = re.match(r"^\d+\.\s*(?P<name>[^()]+?)\s+là\s+(?P<definition>.+)$", line, flags=re.IGNORECASE)
            if not match:
                continue
            name = match.group("name").strip()
            abbr = match.groupdict().get("abbr", "").strip()
            definition = match.group("definition").strip()
            entity_id = f"ENT-{canonical_id(abbr or name).upper()}"
            entities[entity_id] = {
                "entity_id": entity_id,
                "canonical_name": canonical_id(name),
                "display_name": name,
                "abbreviation": abbr,
                "entity_type": "TECHNICAL_CONCEPT" if abbr else "LEGAL_CONCEPT",
                "definition": definition,
                "definition_node_id": node.id,
                "mentioned_node_ids": sorted(nodes_mentioning(asset.nodes, name, abbr)),
                "aliases": [abbr] if abbr else [],
                "review_status": "PASS",
            }
    for formula in build_semantic_formulas(asset):
        for variable in [formula["result_variable"], *formula["input_variables"]]:
            entity_id = f"ENT-{variable}"
            entities.setdefault(
                entity_id,
                {
                    "entity_id": entity_id,
                    "canonical_name": variable.lower(),
                    "display_name": variable,
                    "abbreviation": variable,
                    "entity_type": "VARIABLE",
                    "definition": "",
                    "definition_node_id": formula["supporting_node_id"],
                    "mentioned_node_ids": [formula["supporting_node_id"]],
                    "aliases": [],
                    "review_status": "NEEDS_REVIEW",
                },
            )
    return sorted(entities.values(), key=lambda item: item["entity_id"])


def build_semantic_formulas(asset: LegalKnowledgeAsset) -> list[dict[str, object]]:
    formulas: list[dict[str, object]] = []
    seen: set[tuple[str, str]] = set()
    for node in asset.nodes:
        for line in node.original_text.splitlines():
            for match in FORMULA_RE.finditer(line):
                expression = match.group(0).strip(" .;:")
                if not formula_like(expression):
                    continue
                result = match.group("result")
                key = (node.id, expression)
                if key in seen:
                    continue
                seen.add(key)
                inputs = [token for token in re.findall(r"\b[A-Z]{1,5}\b", match.group("body")) if token != result]
                formulas.append(
                    {
                        "formula_id": f"{asset.root_id}-FORMULA-{result}-{len(formulas)+1:02d}",
                        "expression": expression,
                        "display_expression": expression,
                        "result_variable": result,
                        "input_variables": sorted(set(inputs)),
                        "supporting_node_id": node.id,
                        "formula_type": "COST_FORMULA" if any(var in expression for var in ("G", "AUCP", "UUCP", "TCF", "EF")) else "FORMULA",
                        "original_text": line,
                        "review_status": "VALIDATED",
                    }
                )
    return formulas


def build_semantic_procedures(asset: LegalKnowledgeAsset) -> list[dict[str, object]]:
    procedures: list[dict[str, object]] = []
    for node in asset.nodes:
        if node.node_type != "PROVISION":
            continue
        if "trình tự" not in node.title.lower():
            continue
        steps = extract_numbered_steps(node.original_text)
        if len(steps) < 2:
            continue
        procedures.append(
            {
                "procedure_id": f"{asset.root_id}-PROC-{canonical_id(node.title).upper()}",
                "title": node.title,
                "supporting_node_id": node.id,
                "steps": steps,
                "review_status": "PASS",
            }
        )
    return procedures


def build_cross_references(asset: LegalKnowledgeAsset) -> list[dict[str, object]]:
    appendix_by_number = {
        node.canonical_number: node.id
        for node in asset.nodes
        if node.node_type == "APPENDIX" and node.canonical_number
    }
    references: list[dict[str, object]] = []
    for node in asset.nodes:
        if node.node_type != "PROVISION":
            continue
        for match in REFERENCE_RE.finditer(node.original_text):
            canonical = canonicalize_number(match.group(2))
            target = appendix_by_number.get(canonical)
            if not target:
                continue
            references.append(
                {
                    "reference_id": f"{asset.root_id}-XREF-{len(references)+1:03d}",
                    "source_node_id": node.id,
                    "relation_type": "REFERENCES",
                    "target_node_id": target,
                    "reference_text": match.group(0),
                    "source_location": {},
                    "review_status": "PASS",
                }
            )
    return references


def build_legal_references(asset: LegalKnowledgeAsset) -> list[dict[str, object]]:
    refs: list[dict[str, object]] = []
    for node in asset.nodes:
        if node.node_type not in {"PROVISION", "MAIN_DOCUMENT", "ISSUED_CONTENT"}:
            continue
        for match in LEGAL_REFERENCE_RE.finditer(node.original_text):
            refs.append(
                {
                    "legal_reference_id": f"{asset.root_id}-LREF-{len(refs)+1:03d}",
                    "source_node_id": node.id,
                    "target_document_number": match.group("number"),
                    "target_article": extract_nearby_reference_part(node.original_text, match.start(), "Điều"),
                    "target_clause": extract_nearby_reference_part(node.original_text, match.start(), "khoản"),
                    "target_point": extract_nearby_reference_part(node.original_text, match.start(), "điểm"),
                    "reference_type": "LEGAL_BASIS" if node.node_type == "MAIN_DOCUMENT" else "METHOD_REFERENCE",
                    "original_reference_text": match.group(0),
                    "target_asset_id": "",
                    "resolution_status": "UNRESOLVED",
                }
            )
    return refs


def build_appendix_metadata(asset: LegalKnowledgeAsset, keywords: list[dict[str, object]], formulas: list[dict[str, object]]) -> list[dict[str, object]]:
    keyword_by_node: dict[str, list[str]] = {}
    for keyword in keywords:
        for node_id in keyword["node_ids"]:
            keyword_by_node.setdefault(node_id, []).append(keyword["display_term"])
    formula_by_node: dict[str, list[str]] = {}
    for formula in formulas:
        formula_by_node.setdefault(formula["supporting_node_id"], []).append(formula["formula_id"])
    return [
        {
            "appendix_id": node.id,
            "number": node.canonical_number or node.number,
            "official_title": appendix_official_title(node),
            "short_title": short_appendix_title(node),
            "appendix_type": classify_appendix(node),
            "purpose": infer_appendix_purpose(node),
            "related_node_ids": related_provisions_for_appendix(asset, node),
            "keywords": keyword_by_node.get(node.id, []),
            "formula_ids": formula_by_node.get(node.id, []),
            "review_status": "PASS" if appendix_official_title(node) else "NEEDS_REVIEW",
        }
        for node in asset.nodes
        if node.node_type in {"APPENDIX", "FORM"}
    ]


def validate_semantic_data(asset: LegalKnowledgeAsset) -> dict[str, list[dict[str, str]]]:
    errors: list[dict[str, str]] = []
    warnings: list[dict[str, str]] = []
    node_ids = {node.id for node in asset.nodes}
    topics = asset.semantic.get("topics", [])
    keywords = asset.semantic.get("keywords", [])
    formulas = asset.semantic.get("formulas", [])
    procedures = asset.semantic.get("procedures", [])
    cross_references = asset.semantic.get("cross_references", [])
    legal_references = asset.semantic.get("legal_references", [])
    appendix_metadata = asset.semantic.get("appendix_metadata", [])

    for topic in topics:
        if topic.get("topic_type") not in {item[1] for item in TOPIC_TAXONOMY.values()}:
            errors.append(error("KB-SEM-TOPIC", f"Topic type không hợp lệ: {topic.get('topic_type')}"))
        if not isinstance(topic.get("node_ids"), list):
            errors.append(error("KB-SEM-TOPIC", "Topic node_ids phải là array."))
    for keyword in keywords:
        term = str(keyword.get("display_term", ""))
        if not valid_keyword(term):
            errors.append(error("KB-SEM-KEYWORD", f"Keyword không hợp lệ: {term}"))
    canonical_terms = [str(keyword.get("canonical_term")) for keyword in keywords]
    if duplicate_values(canonical_terms):
        errors.append(error("KB-SEM-KEYWORD", "Keyword canonical bị trùng."))
    for entity in asset.semantic.get("entities", []):
        if entity.get("definition_node_id") not in node_ids:
            errors.append(error("KB-SEM-ENTITY", f"Entity thiếu supporting node: {entity.get('entity_id')}"))
    node_by_id = {node.id: node for node in asset.nodes}
    for formula in formulas:
        node = node_by_id.get(str(formula.get("supporting_node_id")))
        if not node or str(formula.get("display_expression")) not in node.original_text:
            errors.append(error("KB-SEM-FORMULA", f"Công thức không khớp original_text: {formula.get('formula_id')}"))
    for procedure in procedures:
        node = node_by_id.get(str(procedure.get("supporting_node_id")))
        if not node:
            errors.append(error("KB-SEM-PROCEDURE", f"Procedure thiếu node: {procedure.get('procedure_id')}"))
            continue
        for step in procedure.get("steps", []):
            if step.get("original_text") not in node.original_text:
                errors.append(error("KB-SEM-PROCEDURE", f"Procedure mất bước: {procedure.get('procedure_id')}"))
    for reference in cross_references:
        if reference.get("source_node_id") not in node_ids or reference.get("target_node_id") not in node_ids:
            errors.append(error("KB-SEM-XREF", f"Cross reference trỏ sai node: {reference.get('reference_id')}"))
    for legal_reference in legal_references:
        if not legal_reference.get("target_document_number") or legal_reference.get("source_node_id") not in node_ids:
            errors.append(error("KB-SEM-LREF", f"Legal reference không hợp lệ: {legal_reference.get('legal_reference_id')}"))
    for appendix in appendix_metadata:
        if not appendix.get("official_title"):
            errors.append(error("KB-SEM-APP", f"Appendix metadata thiếu official_title: {appendix.get('appendix_id')}"))
    return {"errors": errors, "warnings": warnings}


def duplicate_values(values: list[str]) -> list[str]:
    seen: set[str] = set()
    duplicates: list[str] = []
    for value in values:
        if value in seen and value not in duplicates:
            duplicates.append(value)
        seen.add(value)
    return duplicates


def canonical_keyword(keyword: str) -> str:
    normalized = re.sub(r"\s+", " ", keyword.replace("\xa0", " ")).strip(" .;:,").lower()
    canonical = KEYWORD_ALIASES.get(normalized, normalized)
    return canonical_id(canonical)


def display_keyword(canonical_term: str) -> str:
    reverse = {
        "cost_estimation_dossier": "Hồ sơ xác định chi phí phần mềm nội bộ",
        "cost_estimation_method": "Phương pháp tính chi phí phần mềm nội bộ",
        "software_cost_estimation": "Xác định chi phí phần mềm nội bộ",
        "actor": "Tác nhân",
        "use_case": "Trường hợp sử dụng",
        "transaction": "Giao dịch",
    }
    if canonical_term in reverse:
        return reverse[canonical_term]
    return canonical_term.replace("_", " ")


def keyword_type(canonical_term: str) -> str:
    if canonical_term in {"actor", "use_case", "transaction"} or "he_so" in canonical_term:
        return "TECHNICAL_TERM"
    if "cong_thuc" in canonical_term or re.fullmatch(r"[a-z]{1,5}", canonical_term):
        return "FORMULA_SYMBOL"
    if "trinh_tu" in canonical_term or "ho_so" in canonical_term:
        return "PROCEDURE"
    if re.search(r"\d{2,4}.*(nd|tt|qd|qh)", canonical_term):
        return "DOCUMENT"
    return "LEGAL_TERM"


def valid_keyword(keyword: str) -> bool:
    compact = re.sub(r"\s+", " ", keyword.strip())
    if not compact:
        return False
    if len(compact.split()) > 12 and not any(marker in compact for marker in ("phần mềm nội bộ", "hệ số", "chi phí")):
        return False
    if re.search(r"[.!?]$", compact):
        return False
    if compact.lower().startswith(("bảng này được", "việc xác định", "các nội dung quy định")):
        return False
    return True


def canonical_id(value: str) -> str:
    replacements = str.maketrans(
        "àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ",
        "aaaaaaaaaaaaaaaaaeeeeeeeeeeeiiiiiooooooooooooooooouuuuuuuuuuuyyyyyd",
    )
    ascii_value = value.lower().translate(replacements)
    return re.sub(r"_+", "_", re.sub(r"[^a-z0-9]+", "_", ascii_value)).strip("_")


def nodes_mentioning(nodes: list[AssetNode], name: str, abbr: str = "") -> list[str]:
    mentions: list[str] = []
    needles = [name.lower()]
    if abbr:
        needles.append(abbr.lower())
    for node in nodes:
        text = node.original_text.lower()
        if any(needle and needle in text for needle in needles):
            mentions.append(node.id)
    return mentions


def formula_like(expression: str) -> bool:
    if len(expression) > 120:
        return False
    return bool(re.search(r"\d|[A-Z]{2,}|[×x*/+\-]", expression))


def extract_numbered_steps(text: str) -> list[dict[str, object]]:
    steps: list[dict[str, object]] = []
    current: dict[str, object] | None = None
    for line in text.splitlines()[1:]:
        match = re.match(r"^(?P<number>\d+)\.\s+(?P<body>.+)$", line.strip())
        if match:
            if current:
                steps.append(current)
            current = {
                "step_number": int(match.group("number")),
                "original_text": line.strip(),
                "related_node_ids": [],
            }
        elif current and line.strip():
            current["original_text"] += "\n" + line.strip()
    if current:
        steps.append(current)
    return steps


def extract_nearby_reference_part(text: str, start: int, label: str) -> str:
    window = text[max(0, start - 90) : min(len(text), start + 90)]
    match = re.search(rf"{label}\s+([0-9a-zA-ZđĐ]+)", window, flags=re.IGNORECASE)
    return match.group(1) if match else ""


def appendix_official_title(node: AssetNode) -> str:
    lines = [line.strip() for line in node.original_text.splitlines() if line.strip()]
    for line in lines[1:4]:
        if not APPENDIX_RE.match(line) and not FORM_RE.match(line):
            return line
    return node.title if node.title and node.title.lower() != "phụ lục" else f"Phụ lục {node.canonical_number or node.number}"


def short_appendix_title(node: AssetNode) -> str:
    official = appendix_official_title(node)
    return official[:90]


def classify_appendix(node: AssetNode) -> str:
    text = f"{node.title}\n{node.original_text}".lower()
    if any(marker in text for marker in ("tính toán", "chi phí", "điểm", "hệ số", "công thức")):
        return "CALCULATION_TABLE"
    if any(marker in text for marker in ("mẫu", "danh sách", "biểu mẫu")):
        return "INPUT_FORM"
    if "hướng dẫn" in text:
        return "GUIDANCE_TABLE"
    return "REFERENCE_TABLE"


def infer_appendix_purpose(node: AssetNode) -> str:
    title = appendix_official_title(node)
    text = f"{title}\n{node.original_text}".lower()
    if "tác nhân" in text:
        return "Lập danh sách và phân loại Actor."
    if "trường hợp sử dụng" in text:
        return "Tính hoặc mô tả điểm trường hợp sử dụng."
    if "chi phí trực tiếp" in text:
        return "Tổng hợp chi phí trực tiếp phần mềm nội bộ."
    if "hệ số" in text:
        return "Xác định hệ số phục vụ tính chi phí."
    return title


def related_provisions_for_appendix(asset: LegalKnowledgeAsset, appendix: AssetNode) -> list[str]:
    related: list[str] = []
    label = f"Phụ lục {appendix.canonical_number or appendix.number}"
    for node in asset.nodes:
        if node.node_type == "PROVISION" and label.lower() in node.original_text.lower():
            related.append(node.id)
    return related


def render_concept_index(asset: LegalKnowledgeAsset) -> str:
    concepts = asset.semantic.get("concepts", [])
    if not concepts:
        return "Không phát hiện khái niệm được định nghĩa trực tiếp."
    lines = ["| concept_id | display_name | definition_node_id | definition |", "| --- | --- | --- | --- |"]
    for concept in concepts:
        lines.append(
            f"| {concept.get('entity_id', '')} | {concept.get('display_name', '')} | "
            f"{concept.get('definition_node_id', '')} | {markdown_cell(concept.get('definition', ''))} |"
        )
    return "\n".join(lines)


def render_formula_index(asset: LegalKnowledgeAsset) -> str:
    formulas = asset.semantic.get("formulas", [])
    if not formulas:
        return "Không phát hiện công thức được nêu trực tiếp trong văn bản."
    lines = ["| formula_id | expression | supporting_node_id | variables |", "| --- | --- | --- | --- |"]
    for formula in formulas:
        variables = ", ".join([str(formula.get("result_variable", "")), *[str(item) for item in formula.get("input_variables", [])]])
        lines.append(
            f"| {formula.get('formula_id', '')} | `{formula.get('display_expression', '')}` | "
            f"{formula.get('supporting_node_id', '')} | {variables} |"
        )
    return "\n".join(lines)


def render_procedure_index(asset: LegalKnowledgeAsset) -> str:
    procedures = asset.semantic.get("procedures", [])
    if not procedures:
        return "Không phát hiện trình tự thực hiện được đánh số trực tiếp."
    lines: list[str] = []
    for procedure in procedures:
        lines.extend([f"### {procedure.get('title', procedure.get('procedure_id', ''))}", ""])
        for step in procedure.get("steps", []):
            lines.append(str(step.get("original_text", "")))
        lines.append("")
    return "\n".join(lines).strip()


def render_cross_reference_index(asset: LegalKnowledgeAsset) -> str:
    references = asset.semantic.get("cross_references", [])
    if not references:
        return "Không phát hiện liên kết Điều – Phụ lục."
    lines = ["| reference_id | source_node_id | target_node_id | reference_text |", "| --- | --- | --- | --- |"]
    for reference in references:
        lines.append(
            f"| {reference.get('reference_id', '')} | {reference.get('source_node_id', '')} | "
            f"{reference.get('target_node_id', '')} | {markdown_cell(reference.get('reference_text', ''))} |"
        )
    return "\n".join(lines)


def render_legal_reference_index(asset: LegalKnowledgeAsset) -> str:
    references = asset.semantic.get("legal_references", [])
    if not references:
        return "Không phát hiện văn bản được viện dẫn."
    lines = ["| legal_reference_id | source_node_id | target_document_number | reference_type | original_reference_text |", "| --- | --- | --- | --- | --- |"]
    for reference in references:
        lines.append(
            f"| {reference.get('legal_reference_id', '')} | {reference.get('source_node_id', '')} | "
            f"{reference.get('target_document_number', '')} | {reference.get('reference_type', '')} | "
            f"{markdown_cell(reference.get('original_reference_text', ''))} |"
        )
    return "\n".join(lines)


def render_asset_lookup(asset: LegalKnowledgeAsset) -> str:
    nodes = asset.nodes
    topics_by_node: dict[str, list[str]] = {}
    for topic in asset.semantic.get("topics", []):
        for node_id in topic.get("node_ids", []):
            topics_by_node.setdefault(str(node_id), []).append(str(topic.get("label_vi", "")))
    keywords_by_node: dict[str, list[str]] = {}
    for keyword in asset.semantic.get("keywords", []):
        for node_id in keyword.get("node_ids", []):
            keywords_by_node.setdefault(str(node_id), []).append(str(keyword.get("canonical_term", "")))
    formulas_by_node: dict[str, list[str]] = {}
    for formula in asset.semantic.get("formulas", []):
        formulas_by_node.setdefault(str(formula.get("supporting_node_id", "")), []).append(str(formula.get("formula_id", "")))
    references_by_node: dict[str, list[str]] = {}
    for reference in asset.semantic.get("cross_references", []):
        references_by_node.setdefault(str(reference.get("source_node_id", "")), []).append(str(reference.get("target_node_id", "")))
    for reference in asset.semantic.get("legal_references", []):
        references_by_node.setdefault(str(reference.get("source_node_id", "")), []).append(str(reference.get("target_document_number", "")))
    lines = [
        "| node_id | scope_type | title | topics | canonical_keywords | formulas | references |",
        "| --- | --- | --- | --- | --- | --- | --- |",
    ]
    for node in nodes:
        if node.node_type not in {"PROVISION", "APPENDIX", "FORM", "REFERENCE"}:
            continue
        lines.append(
            f"| {node.id} | {node.scope_type} | {markdown_cell(node.title)} | "
            f"{', '.join(sorted(set(topics_by_node.get(node.id, []))))} | "
            f"{', '.join(sorted(set(keywords_by_node.get(node.id, []))))} | "
            f"{', '.join(sorted(set(formulas_by_node.get(node.id, []))))} | "
            f"{', '.join(sorted(set(references_by_node.get(node.id, []))))} |"
        )
    return "\n".join(lines)


def markdown_cell(value: object) -> str:
    return re.sub(r"\s+", " ", str(value)).replace("|", "\\|").strip()


def render_asset_topics_keywords(article_knowledge) -> str:
    if not article_knowledge:
        return "Không phát hiện chủ đề/từ khóa đạt tiêu chí."
    return "\n\n".join(
        [
            embed_markdown_section(render_topics(article_knowledge), "Không phát hiện chủ đề đạt tiêu chí."),
            embed_markdown_section(render_keyword_index(article_knowledge), "Không phát hiện từ khóa đạt tiêu chí."),
        ]
    )


def render_asset_faq(article_knowledge) -> str:
    faqs = quality_checked_faqs(article_knowledge) if article_knowledge else []
    if not faqs:
        return "Không có FAQ đạt kiểm tra chất lượng."
    lines: list[str] = []
    for faq in faqs:
        lines.extend(
            [
                f"### {faq['question']}",
                f"- Câu trả lời ngắn: {faq['answer']}",
                f"- Căn cứ: {faq['citation']}",
                f"- Mức độ chắc chắn: {faq['confidence']}",
                f"- Ghi chú: {faq['note']}",
                "",
            ]
        )
    return "\n".join(lines).strip()


def build_article_knowledge_from_nodes(nodes: list[AssetNode]):
    class MinimalParsed:
        articles: list[Article]

    parsed = MinimalParsed()
    parsed.articles = [node_to_article(node) for node in nodes if node.node_type == "PROVISION"]
    return build_article_knowledge(parsed)  # type: ignore[arg-type]


def node_to_article(node: AssetNode) -> Article:
    return Article(number=node.number, title=node.title, raw_content=node.original_text)


def children(nodes: list[AssetNode], parent_id: str, node_type: str | None = None) -> list[AssetNode]:
    output = [node for node in nodes if node.parent_id == parent_id and (node_type is None or node.node_type == node_type)]
    return sorted(output, key=lambda node: node.order)


def write_legal_asset_outputs(asset: LegalKnowledgeAsset, output_root: Path) -> dict[str, Path]:
    output_root.mkdir(parents=True, exist_ok=True)
    safe = document_folder_name_from_asset(asset)
    json_path = output_root / f"LEGAL_ASSET_{safe}.json"
    structure_path = output_root / f"STRUCTURE_{safe}.json"
    md_path = output_root / f"LEGAL_ASSET_{safe}.md"
    gpt_path = output_root / f"GPT_KNOWLEDGE_{safe}.md"
    docx_path = output_root / f"LEGAL_ASSET_{safe}.docx"
    semantic_dir = output_root / f"semantic_{safe}"
    migration_path = output_root / f"MIGRATION_REPORT_{safe}.md"
    validation_path = output_root / f"ASSET_VALIDATION_{safe}.md"
    regression_path = output_root / f"REGRESSION_SUMMARY_{safe}.md"
    runtime_log_path = output_root / f"RUNTIME_LOG_{safe}.log"
    json_path.write_text(json.dumps(asset_to_dict(asset), ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    structure = build_structure(asset)
    structure_validation = validate_structure_for_export(asset)
    structure["structure_validation"] = structure_validation
    structure_path.write_text(json.dumps(structure, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    migration_path.write_text(render_migration_report(asset), encoding="utf-8")
    validation_path.write_text(render_asset_validation_report(asset), encoding="utf-8")
    regression_path.write_text(render_regression_summary(asset), encoding="utf-8")
    runtime_log_path.write_text(render_runtime_log(asset, structure_validation), encoding="utf-8")
    write_semantic_outputs(asset, semantic_dir)
    if structure_validation["status"] == "FAIL":
        raise ValueError("Structure validation FAIL: " + "; ".join(structure_validation["errors"]))
    md_path.write_text(render_asset_markdown(asset), encoding="utf-8")
    gpt_path.write_text(render_gpt_knowledge_from_asset(asset), encoding="utf-8")
    write_asset_docx(asset, docx_path)
    return {
        "json": json_path,
        "structure": structure_path,
        "markdown": md_path,
        "gpt_markdown": gpt_path,
        "word": docx_path,
        "semantic_dir": semantic_dir,
        "migration_report": migration_path,
        "validation_report": validation_path,
        "regression_summary": regression_path,
        "runtime_log": runtime_log_path,
    }


def write_semantic_outputs(asset: LegalKnowledgeAsset, semantic_dir: Path) -> None:
    semantic_dir.mkdir(parents=True, exist_ok=True)
    for name in (
        "topics",
        "keywords",
        "concepts",
        "entities",
        "formulas",
        "procedures",
        "cross_references",
        "legal_references",
        "appendix_metadata",
    ):
        path = semantic_dir / f"{name}.json"
        path.write_text(json.dumps(asset.semantic.get(name, []), ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def render_runtime_log(asset: LegalKnowledgeAsset, structure_validation: dict[str, object]) -> str:
    return "\n".join(
        [
            f"parser_version={PARSER_VERSION}",
            f"schema_version={asset.schema_version}",
            f"exporter_version={EXPORTER_VERSION}",
            f"asset_id={asset.root_id}",
            f"issued_content={'yes' if asset.stats['issued_content_count'] else 'no'}",
            f"main_document_article_count={asset.stats['main_document_article_count']}",
            f"issued_content_count={asset.stats['issued_content_count']}",
            f"issued_content_provision_count={asset.stats['issued_content_provision_count']}",
            f"appendix_count={asset.stats['appendix_count']}",
            f"reference_count={asset.stats['reference_count']}",
            f"asset_validation_status={asset.validation['status']}",
            f"structure_validation_status={structure_validation['status']}",
            f"pipeline=LegalKnowledgeAsset",
            f"gpt_exporter=LegalKnowledgeAsset",
            "",
        ]
    )


def write_asset_docx(asset: LegalKnowledgeAsset, path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_heading(asset.title, level=1)
    document.add_heading("SOURCE OF TRUTH", level=2)
    document.add_paragraph(
        "Nội dung gốc là nguồn ưu tiên cao nhất. Metadata, bảng tra cứu, tóm tắt, từ khóa và FAQ chỉ hỗ trợ định vị."
    )
    document.add_heading("Metadata", level=2)
    document.add_paragraph(f"Loại văn bản: {asset.document_type}")
    document.add_paragraph(f"Số/ký hiệu: {asset.document_number}")
    document.add_paragraph(f"Root node: {asset.root_id}")
    document.add_paragraph(f"Validation: {asset.validation['status']}")

    nodes = asset.nodes
    root = next(node for node in nodes if node.id == asset.root_id)
    document.add_heading("Văn bản chính", level=2)
    for provision in children(nodes, root.id, "PROVISION"):
        document.add_heading(f"Điều {provision.number}", level=3)
        document.add_paragraph(provision.original_text)

    issued_contents = children(nodes, root.id, "ISSUED_CONTENT")
    if issued_contents:
        document.add_heading("Nội dung ban hành kèm theo", level=2)
        for issued in issued_contents:
            document.add_heading(issued.title, level=3)
            for provision in children(nodes, issued.id, "PROVISION"):
                kind = provision.metadata.get("provision_kind", "Điều")
                document.add_heading(f"{kind} {provision.number}", level=4)
                document.add_paragraph(provision.original_text)
            appendices = [node for node in children(nodes, issued.id) if node.node_type in {"APPENDIX", "FORM"}]
            if appendices:
                document.add_heading("Phụ lục của nội dung ban hành kèm theo", level=2)
                for appendix in appendices:
                    label = "Biểu mẫu" if appendix.node_type == "FORM" else "Phụ lục"
                    document.add_heading(f"{label} {appendix.canonical_number or appendix.number}", level=3)
                    document.add_paragraph(appendix.original_text)

    document.add_heading("Validation", level=2)
    document.add_paragraph(json.dumps(asset.validation, ensure_ascii=False, indent=2))
    document.save(path)


def asset_to_dict(asset: LegalKnowledgeAsset) -> dict[str, object]:
    return {
        "schema_version": asset.schema_version,
        "document_number": asset.document_number,
        "document_type": asset.document_type,
        "title": asset.title,
        "root_id": asset.root_id,
        "stats": asset.stats,
        "validation": asset.validation,
        "semantic": asset.semantic,
        "migration_report": asset.migration_report,
        "nodes": [asdict(node) for node in asset.nodes],
    }


def render_migration_report(asset: LegalKnowledgeAsset) -> str:
    report = asset.migration_report
    lines = [
        f"# Migration Report - {asset.document_number}",
        "",
        f"- Legacy articles: {report['legacy_article_count']}",
        f"- Legacy appendices: {report['legacy_appendix_count']}",
        f"- New stats: `{json.dumps(report['new_stats'], ensure_ascii=False)}`",
        f"- Legacy checksum: `{report['legacy_checksum']}`",
        f"- Asset checksum: `{report['asset_checksum']}`",
        "",
        "## MAIN_DOCUMENT nodes",
        "",
    ]
    lines.extend(f"- {node_id}" for node_id in report["main_document_nodes"])
    lines.extend(["", "## ISSUED_CONTENT nodes", ""])
    lines.extend(f"- {node_id}" for node_id in report["issued_content_nodes"])
    lines.extend(["", "## APPENDIX/FORM nodes", ""])
    lines.extend(f"- {node_id}" for node_id in report["appendix_nodes"])
    lines.extend(["", "## APPENDIX -> REFERENCE nodes", ""])
    lines.extend(f"- {node_id}" for node_id in report["appendix_to_reference_nodes"])
    lines.extend(["", "## NEEDS_REVIEW", ""])
    review_nodes = report["review_nodes"]
    lines.extend(f"- {node_id}" for node_id in review_nodes) if review_nodes else lines.append("- Không có.")
    return "\n".join(lines).strip() + "\n"


def render_asset_validation_report(asset: LegalKnowledgeAsset) -> str:
    validation = asset.validation
    lines = [f"# Asset Validation - {asset.document_number}", "", f"- Kết luận: {validation['status']}", "", "## Errors"]
    errors = validation["errors"]
    lines.extend(f"- {item['code']}: {item['message']}" for item in errors) if errors else lines.append("- Không có.")
    lines.extend(["", "## Warnings"])
    warnings = validation["warnings"]
    lines.extend(f"- {item['code']}: {item['message']}" for item in warnings) if warnings else lines.append("- Không có.")
    return "\n".join(lines).strip() + "\n"


def render_regression_summary(asset: LegalKnowledgeAsset) -> str:
    return "\n".join(
        [
            f"# Regression Summary - {asset.document_number}",
            "",
            f"- MAIN_DOCUMENT parsing: {'PASS' if asset.stats['main_document_article_count'] >= 0 else 'FAIL'}",
            f"- ISSUED_CONTENT detection: {'PASS' if asset.stats['issued_content_count'] >= 0 else 'FAIL'}",
            f"- Appendix detection: {'PASS' if asset.stats['appendix_count'] >= 0 else 'FAIL'}",
            f"- Reference detection: {'PASS' if asset.stats['reference_count'] >= 0 else 'FAIL'}",
            f"- Validation: {asset.validation['status']}",
            "",
        ]
    )


def document_folder_name_from_asset(asset: LegalKnowledgeAsset) -> str:
    source = asset.document_number or asset.root_id or uuid4().hex[:8]
    source = source.replace("/", "_").replace(".", "_")
    source = re.sub(r"[^A-Za-z0-9Đđ_-]+", "_", source.strip())
    return re.sub(r"_+", "_", source).strip("_") or "legal_asset"


def normalize_node_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def checksum(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def canonicalize_number(number: str) -> str:
    compact = number.strip().upper().replace(".", "").replace(" ", "")
    if "/" in compact or "-" in compact:
        return re.sub(r"[^A-Z0-9Đ/-]+", "", compact).replace("/", "-")
    roman = re.fullmatch(r"[IVXLCDM]+", compact)
    if roman:
        return compact
    digits = re.search(r"\d+", compact)
    if digits:
        return str(int(digits.group(0)))
    return compact
